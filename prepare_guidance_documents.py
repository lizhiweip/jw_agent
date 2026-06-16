"""Extract normative case-handling documents into resumable JSON records.

Run this script with the dedicated OCR environment. Text PDFs are extracted
directly; scanned PDF pages are recognized with PaddleOCR. Page results are
cached individually so a long OCR run can safely resume after interruption.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Iterable


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xlsx", ".txt"}
TEXT_PAGE_THRESHOLD = 80
CHUNK_SIZE = 900
TARGET_CHUNK_SIZE = 650
CHUNK_OVERLAP = 100

CHINESE_NUMBER = "一二三四五六七八九十百千万零〇两"
STRUCTURE_PATTERNS = (
    ("part", re.compile(rf"^第[{CHINESE_NUMBER}\d]+编")),
    ("chapter", re.compile(rf"^第[{CHINESE_NUMBER}\d]+章")),
    ("section", re.compile(rf"^第[{CHINESE_NUMBER}\d]+节")),
    ("article", re.compile(rf"^第[{CHINESE_NUMBER}\d]+条")),
    ("cn_item", re.compile(rf"^[{CHINESE_NUMBER}]+[、．.]\s*")),
    ("number_item", re.compile(r"^\d{1,3}[、．.]\s*")),
    ("subitem", re.compile(rf"^[（(][{CHINESE_NUMBER}\d]+[）)]\s*")),
)
HEADING_KINDS = {"part", "chapter", "section"}
CONTENT_ANCHOR_KINDS = {"article", "cn_item", "number_item", "subitem"}


def clean_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]*\n+", "\n", text)
    text = re.sub(r"(?m)^\s*[-—]?\s*\d+\s*[-—]?\s*$", "", text)
    return text.strip()


def stable_id(relative_path: str) -> str:
    return hashlib.sha1(relative_path.encode("utf-8")).hexdigest()[:16]


def extract_document_number(text: str, fallback_name: str) -> str:
    candidates = "\n".join((fallback_name, text[:2500]))
    patterns = (
        r"[\u4e00-\u9fff]{1,12}(?:办|厅|发|函)?〔\d{4}〕\d+号",
        r"[\u4e00-\u9fff]{1,12}(?:办|厅|发|函)?\[\d{4}\]\d+号",
        r"[\u4e00-\u9fff]{1,12}(?:办|厅|发|函)?【\d{4}】\d+号",
    )
    for pattern in patterns:
        match = re.search(pattern, candidates)
        if match:
            return match.group(0)
    return ""


def title_from_filename(path: Path) -> str:
    title = path.stem
    title = re.sub(r"^\d+(?:\.\d+)*[.、．\s]*", "", title)
    title = re.sub(r"\s*\[\d+\]\s*$", "", title)
    return title.strip() or path.stem


def structure_kind(text: str) -> str:
    compact = text.strip()
    for kind, pattern in STRUCTURE_PATTERNS:
        if pattern.match(compact):
            return kind
    return ""


def normalize_ocr_page(page_text: str) -> list[str]:
    """Repair visual line wrapping while preserving legal document structure."""
    raw_lines = [
        clean_text(line)
        for line in page_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        if clean_text(line)
    ]
    paragraphs: list[str] = []
    current = ""

    for line in raw_lines:
        kind = structure_kind(line)
        if kind:
            if current:
                paragraphs.append(current)
            current = line
            continue

        if not current:
            current = line
            continue

        # OCR usually inserts a newline at the visual right margin. Chinese text
        # should be rejoined without a space unless the previous line clearly
        # completed a paragraph.
        if re.search(r"[。！？；：]$|[）)]$|[》〉]$", current):
            paragraphs.append(current)
            current = line
        else:
            current += line

    if current:
        paragraphs.append(current)
    return [clean_text(item) for item in paragraphs if clean_text(item)]


def split_sentences(text: str) -> list[str]:
    sentences = [
        item.strip()
        for item in re.split(r"(?<=[。！？；])", clean_text(text))
        if item.strip()
    ]
    return sentences or ([clean_text(text)] if clean_text(text) else [])


def hard_split_sentence(text: str, size: int) -> list[str]:
    """Last-resort split for OCR text without sentence punctuation."""
    pieces: list[str] = []
    remaining = text.strip()
    while len(remaining) > size:
        cut = max(
            remaining.rfind(mark, 0, size + 1)
            for mark in ("，", "、", "：", ",", " ")
        )
        if cut < size // 2:
            cut = size
        else:
            cut += 1
        pieces.append(remaining[:cut].strip())
        remaining = remaining[cut:].strip()
    if remaining:
        pieces.append(remaining)
    return pieces


def split_semantic_section(section: dict, size: int = CHUNK_SIZE) -> list[dict]:
    text = clean_text(section["content"])
    if len(text) <= size:
        return [{**section, "content": text}]

    sentences: list[str] = []
    for sentence in split_sentences(text):
        sentences.extend(
            hard_split_sentence(sentence, size)
            if len(sentence) > size
            else [sentence]
        )

    pieces: list[dict] = []
    current: list[str] = []
    for sentence in sentences:
        candidate = "".join(current) + sentence
        if current and len(candidate) > size:
            pieces.append({**section, "content": "".join(current)})
            overlap = current[-1] if len(current[-1]) <= CHUNK_OVERLAP else ""
            current = [overlap, sentence] if overlap else [sentence]
        else:
            current.append(sentence)
    if current:
        pieces.append({**section, "content": "".join(current)})
    return pieces


def build_semantic_sections(
    pages: list[tuple[int | None, str]],
) -> list[dict]:
    sections: list[dict] = []
    current: dict | None = None
    heading_path: list[str] = []
    topic_heading = ""
    active_anchor = ""
    pending_headings: list[tuple[str, int | None]] = []

    def flush() -> None:
        nonlocal current
        if current and clean_text(current["content"]):
            current["content"] = clean_text(current["content"])
            sections.append(current)
        current = None

    for page_number, page_text in pages:
        for paragraph in normalize_ocr_page(page_text):
            kind = structure_kind(paragraph)
            if kind in HEADING_KINDS:
                flush()
                level = {"part": 0, "chapter": 1, "section": 2}[kind]
                heading_path = heading_path[:level]
                heading_path.append(paragraph)
                topic_heading = ""
                active_anchor = ""
                pending_headings.append((paragraph, page_number))
                continue

            if kind == "cn_item" and len(paragraph) <= 40:
                flush()
                topic_heading = paragraph
                active_anchor = ""
                pending_headings.append((paragraph, page_number))
                continue

            if kind in CONTENT_ANCHOR_KINDS:
                flush()
                prefix = "\n".join(item[0] for item in pending_headings)
                start_page = (
                    pending_headings[0][1] if pending_headings else page_number
                )
                pending_headings = []
                inherited_path = list(heading_path)
                if topic_heading and topic_heading not in inherited_path:
                    inherited_path.append(topic_heading)
                if (
                    kind == "subitem"
                    and active_anchor
                    and active_anchor not in inherited_path
                ):
                    inherited_path.append(active_anchor)
                current = {
                    "content": f"{prefix}\n{paragraph}".strip(),
                    "page_start": start_page,
                    "page_end": page_number,
                    "section_title": paragraph[:120],
                    "heading_path": inherited_path,
                    "structure_kind": kind,
                }
                if kind in {"article", "number_item", "cn_item"}:
                    active_anchor = paragraph[:120]
                continue

            if current is None:
                prefix = "\n".join(item[0] for item in pending_headings)
                start_page = (
                    pending_headings[0][1] if pending_headings else page_number
                )
                pending_headings = []
                current = {
                    "content": f"{prefix}\n{paragraph}".strip(),
                    "page_start": start_page,
                    "page_end": page_number,
                    "section_title": "",
                    "heading_path": list(heading_path)
                    + ([topic_heading] if topic_heading else []),
                    "structure_kind": "preamble",
                }
            else:
                separator = (
                    "\n"
                    if re.search(r"[。！？；：]$|[）)]$|[》〉]$", current["content"])
                    else ""
                )
                current["content"] += separator + paragraph
                current["page_end"] = page_number

    flush()
    if pending_headings:
        sections.append(
            {
                "content": "\n".join(item[0] for item in pending_headings),
                "page_start": pending_headings[0][1],
                "page_end": pending_headings[-1][1],
                "section_title": pending_headings[-1][0],
                "heading_path": list(heading_path),
                "structure_kind": "heading",
            }
        )
    return sections


def build_page_chunks(pages: list[tuple[int | None, str]]) -> list[dict]:
    """Build structure-aware chunks and pack only adjacent sibling clauses."""
    split_sections: list[dict] = []
    for section in build_semantic_sections(pages):
        split_sections.extend(split_semantic_section(section))

    chunks: list[dict] = []
    for section in split_sections:
        content = clean_text(section["content"])
        previous = chunks[-1] if chunks else None
        same_context = (
            previous is not None
            and previous.get("heading_path", []) == section.get("heading_path", [])
        )
        both_preamble = (
            previous is not None
            and previous["structure_kind"] in {"preamble", "heading"}
            and section["structure_kind"] in {"preamble", "heading"}
        )
        both_structured = (
            previous is not None
            and previous["structure_kind"] not in {"preamble", "heading"}
            and section["structure_kind"] not in {"preamble", "heading"}
        )
        can_merge = bool(
            previous
            and same_context
            and (both_preamble or both_structured)
            and len(previous["content"]) < TARGET_CHUNK_SIZE
            and len(previous["content"]) + len(content) + 1 <= CHUNK_SIZE
        )
        if can_merge:
            previous["content"] += "\n" + content
            previous["page_end"] = section["page_end"]
            titles = previous.setdefault(
                "section_titles",
                [previous["section_title"]] if previous["section_title"] else [],
            )
            if section["section_title"] and section["section_title"] not in titles:
                titles.append(section["section_title"])
            if len(titles) > 1:
                previous["section_title"] = f"{titles[0]} 至 {titles[-1]}"
            previous["structure_kind"] = "group"
        else:
            chunks.append(
                {
                    **section,
                    "content": content,
                    "section_titles": (
                        [section["section_title"]]
                        if section["section_title"]
                        else []
                    ),
                }
            )
    return chunks


class PdfExtractor:
    def __init__(
        self,
        cache_root: Path,
        force_ocr: bool = False,
        max_pages_per_pdf: int = 0,
    ):
        self.cache_root = cache_root
        self.force_ocr = force_ocr
        self.max_pages_per_pdf = max_pages_per_pdf
        self._ocr = None

    def _get_ocr(self):
        if self._ocr is None:
            os.environ.setdefault("PADDLE_PDX_MODEL_SOURCE", "BOS")
            os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
            from paddleocr import PaddleOCR

            self._ocr = PaddleOCR(
                lang="ch",
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=False,
            )
        return self._ocr

    @staticmethod
    def _ocr_result_text(result) -> str:
        data = result.json if hasattr(result, "json") else result
        if isinstance(data, dict):
            payload = data.get("res", data)
            texts = payload.get("rec_texts", [])
            scores = payload.get("rec_scores", [])
            return "\n".join(
                text.strip()
                for text, score in zip(texts, scores)
                if text.strip() and float(score) >= 0.45
            )
        return ""

    def _ocr_page(self, page) -> str:
        import fitz
        import numpy as np

        pix = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
        image = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, pix.n
        )
        results = list(self._get_ocr().predict(image))
        return clean_text("\n".join(self._ocr_result_text(item) for item in results))

    def extract(self, path: Path, document_id: str) -> list[tuple[int, str]]:
        import fitz

        doc = fitz.open(path)
        document_cache = self.cache_root / document_id
        document_cache.mkdir(parents=True, exist_ok=True)
        pages: list[tuple[int, str]] = []

        page_count = (
            min(len(doc), self.max_pages_per_pdf)
            if self.max_pages_per_pdf
            else len(doc)
        )
        for index in range(page_count):
            page = doc[index]
            page_number = index + 1
            cache_file = document_cache / f"{page_number:05d}.json"
            if cache_file.exists():
                cached = json.loads(cache_file.read_text(encoding="utf-8"))
                pages.append((page_number, cached.get("text", "")))
                continue

            direct_text = clean_text(page.get_text("text"))
            method = "text"
            if self.force_ocr or len(direct_text) < TEXT_PAGE_THRESHOLD:
                method = "ocr"
                direct_text = self._ocr_page(page)

            cache_file.write_text(
                json.dumps(
                    {"page": page_number, "method": method, "text": direct_text},
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            pages.append((page_number, direct_text))
            print(
                f"    page {page_number}/{len(doc)}: {method}, "
                f"{len(direct_text)} chars",
                flush=True,
            )
        return pages


def extract_docx(path: Path) -> list[tuple[None, str]]:
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                parts.append(" | ".join(cells))
    return [(None, clean_text("\n".join(parts)))]


def convert_doc_to_docx(path: Path, target: Path) -> Path:
    if os.name != "nt":
        raise RuntimeError("Linux 上请先使用 LibreOffice 将 .doc 转成 .docx")
    import win32com.client

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    document = None
    try:
        document = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        document.SaveAs2(str(target.resolve()), FileFormat=16)
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        try:
            word.Quit()
        except Exception:
            # Word occasionally reports an RPC disconnect after SaveAs2 even
            # though the converted file was written successfully.
            pass
    if not target.exists() or target.stat().st_size == 0:
        raise RuntimeError(f"Word 未能转换文件：{path}")
    return target


def extract_xlsx(path: Path) -> list[tuple[None, str]]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"【工作表】{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [clean_text(str(value)) for value in row if value not in (None, "")]
            if values:
                parts.append(" | ".join(values))
    return [(None, clean_text("\n".join(parts)))]


def extract_text(path: Path) -> list[tuple[None, str]]:
    for encoding in ("utf-8", "gb18030"):
        try:
            return [(None, clean_text(path.read_text(encoding=encoding)))]
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法识别编码: {path}")


def iter_source_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if (
            path.is_file()
            and path.suffix.lower() in SUPPORTED_EXTENSIONS
            and not path.name.startswith("~$")
            and path.stat().st_size > 0
        ):
            yield path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, help="规范性文件源目录")
    parser.add_argument(
        "--output",
        default="knowledge_base_guidance",
        help="结构化 JSON 输出目录",
    )
    parser.add_argument("--limit", type=int, default=0, help="仅处理前 N 个文件")
    parser.add_argument("--contains", default="", help="仅处理路径中包含该文字的文件")
    parser.add_argument(
        "--max-pages-per-pdf",
        type=int,
        default=0,
        help="每份 PDF 最多处理 N 页；0 表示全部",
    )
    parser.add_argument("--force-ocr", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source_root = Path(args.source).resolve()
    output_root = Path(args.output).resolve()
    records_dir = output_root / "records"
    cache_root = output_root / "page_cache"
    records_dir.mkdir(parents=True, exist_ok=True)
    cache_root.mkdir(parents=True, exist_ok=True)

    files = list(iter_source_files(source_root))
    if args.contains:
        files = [
            path
            for path in files
            if args.contains.lower() in str(path.relative_to(source_root)).lower()
        ]
    if args.limit:
        files = files[: args.limit]
    print(f"Found {len(files)} supported files under {source_root}")

    pdf_extractor = PdfExtractor(
        cache_root,
        force_ocr=args.force_ocr,
        max_pages_per_pdf=args.max_pages_per_pdf,
    )
    manifest: list[dict] = []
    failures: list[dict] = []

    for file_index, path in enumerate(files, 1):
        relative_path = path.relative_to(source_root).as_posix()
        document_id = stable_id(relative_path)
        print(f"[{file_index}/{len(files)}] {relative_path}", flush=True)
        try:
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                pages = pdf_extractor.extract(path, document_id)
            elif suffix == ".docx":
                pages = extract_docx(path)
            elif suffix == ".doc":
                with tempfile.TemporaryDirectory() as temp_dir:
                    converted = Path(temp_dir) / f"{path.stem}.docx"
                    pages = extract_docx(convert_doc_to_docx(path, converted))
            elif suffix == ".xlsx":
                pages = extract_xlsx(path)
            elif suffix == ".txt":
                pages = extract_text(path)
            else:
                continue

            chunks = build_page_chunks(pages)
            combined_preview = "\n".join(item[1] for item in pages[:5])
            title = title_from_filename(path)
            document_number = extract_document_number(combined_preview, path.name)
            confidentiality = (
                "内部"
                if "内部" in combined_preview[:300] or "内部" in path.name
                else ""
            )

            for old_record in records_dir.glob(f"guidance_{document_id}_*.json"):
                old_record.unlink()
            for chunk_index, chunk in enumerate(chunks, 1):
                record_id = f"guidance_{document_id}_{chunk_index:04d}"
                record = {
                    "pid": record_id,
                    "document_id": document_id,
                    "knowledge_type": "办案规范",
                    "title": title,
                    "document_number": document_number,
                    "confidentiality": confidentiality,
                    "source_file": relative_path,
                    "page_start": chunk["page_start"],
                    "page_end": chunk["page_end"],
                    "chunk_index": chunk_index,
                    "section_title": chunk.get("section_title", ""),
                    "section_titles": chunk.get("section_titles", []),
                    "heading_path": chunk.get("heading_path", []),
                    "structure_kind": chunk.get("structure_kind", ""),
                    "content": chunk["content"],
                }
                (records_dir / f"{record_id}.json").write_text(
                    json.dumps(record, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )

            manifest.append(
                {
                    "document_id": document_id,
                    "source_file": relative_path,
                    "chunks": len(chunks),
                    "status": "ok",
                }
            )
            print(f"  -> {len(chunks)} chunks", flush=True)
        except Exception as exc:
            failure = {"source_file": relative_path, "error": str(exc)}
            failures.append(failure)
            manifest.append({**failure, "status": "failed"})
            print(f"  FAILED: {exc}", file=sys.stderr, flush=True)

        (output_root / "manifest.json").write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    print(f"Completed: {len(manifest) - len(failures)} ok, {len(failures)} failed")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
