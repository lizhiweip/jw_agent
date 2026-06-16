"""Convert numbered DOCX question/answer pairs into structured JSON records."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


PAIR_PATTERN = re.compile(r"^(?P<number>\d+)(?P<answer>\.1)?\s*$")


def clean_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(path: Path) -> str:
    document = Document(path)
    parts = [
        clean_text(paragraph.text)
        for paragraph in document.paragraphs
        if clean_text(paragraph.text)
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def infer_category(question: str) -> str:
    match = re.search(r"【([^】]+)】", question)
    return match.group(1).strip() if match else "纪检监察业务题"


def collect_pairs(source: Path) -> dict[int, dict[str, Path]]:
    pairs: dict[int, dict[str, Path]] = {}
    for path in sorted(source.glob("*.docx")):
        match = PAIR_PATTERN.fullmatch(path.stem)
        if not match:
            continue
        number = int(match.group("number"))
        role = "answer" if match.group("answer") else "question"
        pairs.setdefault(number, {})[role] = path
    return pairs


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, help="编号 DOCX 所在目录")
    parser.add_argument(
        "--output",
        default="knowledge_base_qa",
        help="结构化问答记录输出目录",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    records_dir = output / "records"
    records_dir.mkdir(parents=True, exist_ok=True)

    pairs = collect_pairs(source)
    incomplete = {
        number: pair
        for number, pair in pairs.items()
        if "question" not in pair or "answer" not in pair
    }
    if incomplete:
        details = ", ".join(str(number) for number in sorted(incomplete))
        raise SystemExit(f"以下编号缺少题目或答案文件：{details}")

    for old_record in records_dir.glob("qa_*.json"):
        old_record.unlink()

    manifest: list[dict] = []
    for number in sorted(pairs):
        question_path = pairs[number]["question"]
        answer_path = pairs[number]["answer"]
        question = extract_docx(question_path)
        answer = extract_docx(answer_path)
        if not question or not answer:
            raise SystemExit(f"第 {number} 组题目或答案为空")

        pid = f"qa_{number:03d}"
        record = {
            "pid": pid,
            "knowledge_type": "业务问答",
            "number": number,
            "title": f"纪检监察业务问答第{number}题",
            "category": infer_category(question),
            "question": question,
            "answer": answer,
            "question_file": question_path.name,
            "answer_file": answer_path.name,
        }
        (records_dir / f"{pid}.json").write_text(
            json.dumps(record, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        manifest.append(
            {
                "pid": pid,
                "number": number,
                "category": record["category"],
                "question_file": question_path.name,
                "answer_file": answer_path.name,
                "question_chars": len(question),
                "answer_chars": len(answer),
                "status": "ok",
            }
        )

    (output / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"业务问答记录生成完成：{len(manifest)} 组，输出到 {records_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
